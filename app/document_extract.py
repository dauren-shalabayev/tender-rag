"""Извлечение текста из PDF и DOCX для индексации."""

from __future__ import annotations

import io
import re

from pypdf import PdfReader


def _norm(s: str) -> str:
    t = s.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _norm("\n".join(parts))


def extract_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _norm("\n".join(parts))


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if name.endswith(".docx"):
        return extract_text_from_docx(data)
    raise ValueError("Поддерживаются только .pdf и .docx")
